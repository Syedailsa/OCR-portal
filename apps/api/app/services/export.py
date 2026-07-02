from io import BytesIO

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def build_docx_bytes(title: str, text: str | None) -> bytes:
    buffer = BytesIO()
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    doc.add_paragraph(text or "")
    doc.save(buffer)
    return buffer.getvalue()


def build_pdf_bytes(title: str, text: str | None) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    pdf.setTitle(title)
    pdf.drawString(40, height - 50, title)
    y = height - 80
    for line in (text or "").splitlines() or [""]:
        pdf.drawString(40, y, line[:120])
        y -= 18
        if y < 60:
            pdf.showPage()
            y = height - 50
    pdf.save()
    return buffer.getvalue()
